"""Generate comprehensive DOCX report for eMenu Smart Tags application."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ─── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(1)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_table_row(table, cells_data, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(text)
        if bold:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('eMenu Smart Tags', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('AI-Powered Guest Intelligence Platform', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Technical Documentation & Architecture Report\n').bold = True
meta.add_run('\nVersion 3.0.0\nFebruary 2026\n\n')
meta.add_run('Live URL: ').bold = True
meta.add_run('https://emenu-smart-tags-ui.onrender.com\n')
meta.add_run('API URL: ').bold = True
meta.add_run('https://smart-tags-predictor.onrender.com\n')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Technology Stack',
    '3. System Architecture',
    '4. Machine Learning Model',
    '   4.1 ANN Architecture',
    '   4.2 Domain Adaptation Layer',
    '   4.3 Blended Scoring System',
    '   4.4 Heuristic Reliability Engine',
    '   4.5 Confidence Calculation',
    '5. Backend API',
    '   5.1 Endpoints',
    '   5.2 Request/Response Schemas',
    '   5.3 Smart Tag Extraction',
    '   5.4 Sentiment Analysis',
    '   5.5 Demo Scenarios',
    '6. Frontend Application',
    '   6.1 Pages Overview',
    '   6.2 Dashboard Page',
    '   6.3 Analyze Page',
    '   6.4 Table View Page',
    '   6.5 History Page',
    '   6.6 Settings Page',
    '7. UI Component Library',
    '   7.1 RiskGauge',
    '   7.2 SmartTagBadge System',
    '   7.3 SmartActions',
    '   7.4 NumberStepper & ChannelSelector',
    '   7.5 SimulatorControls',
    '   7.6 VoiceCommand',
    '   7.7 GuestInsightCard',
    '   7.8 GuestDetailView',
    '8. Data Flow & Prediction Pipeline',
    '9. Styling & Design System',
    '10. Deployment & Infrastructure',
    '11. File Structure',
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'eMenu Smart Tags is a SaaS-grade restaurant intelligence platform that predicts guest no-show risk '
    'using an Artificial Neural Network (ANN) originally trained on hotel reservation data. The system '
    'applies a Domain Adaptation Layer to bridge the gap between hotel and restaurant contexts, then '
    'blends the ANN output with a restaurant-tuned heuristic to produce calibrated risk scores.'
)
doc.add_paragraph(
    'The platform provides restaurant operators with:'
)
add_bullet('Real-time no-show risk prediction (0-100% scale)')
add_bullet('Smart tag extraction from reservation notes (dietary, occasion, seating, accessibility)')
add_bullet('Sentiment analysis of guest notes using TextBlob NLP')
add_bullet('Actionable recommendations based on prediction results')
add_bullet('Voice command input using the Web Speech API')
add_bullet('Time-Travel Simulator for what-if analysis')
add_bullet('Batch prediction for tonight\'s full reservation list')
add_bullet('Analysis history with localStorage persistence')
add_bullet('18 comprehensive demo scenarios for demonstration')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Technology Stack', level=1)

doc.add_heading('Frontend', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Technology'
hdr[1].text = 'Version'
hdr[2].text = 'Purpose'
for r in hdr:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

fe_stack = [
    ('React', '18.3.1', 'UI component library (functional components with hooks)'),
    ('TypeScript', '5.6.2', 'Type-safe JavaScript for compile-time error checking'),
    ('Vite', '6.0.0', 'Build tool and dev server with HMR (Hot Module Replacement)'),
    ('Tailwind CSS', '4.0.0', 'Utility-first CSS framework with JIT compilation'),
    ('Framer Motion', '12.34.0', 'Animation library for spring physics, layout animations, gestures'),
    ('React Router DOM', '7.1.0', 'Client-side routing with nested layouts'),
    ('Recharts', '3.7.0', 'Chart library (available for dashboard extensions)'),
    ('Lucide React', '0.468.0', 'Icon library (200+ icons used across components)'),
    ('Vaul', '1.1.2', 'Mobile-first drawer component (bottom sheet on mobile)'),
]
for row_data in fe_stack:
    add_table_row(table, row_data)

doc.add_paragraph()
doc.add_heading('Backend', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Technology'
hdr[1].text = 'Version'
hdr[2].text = 'Purpose'
for r in hdr:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

be_stack = [
    ('Python', '3.11+', 'Backend runtime'),
    ('FastAPI', '0.115.0+', 'Async web framework with automatic OpenAPI docs'),
    ('Pydantic', '2.9.0+', 'Data validation and serialization for API schemas'),
    ('TensorFlow / Keras', '2.15.0+', 'Deep learning framework for ANN model inference'),
    ('scikit-learn', '1.4.0+', 'StandardScaler and OneHotEncoder for feature preprocessing'),
    ('Pandas', '2.1.0+', 'DataFrame operations for feature engineering and data simulation'),
    ('NumPy', '1.24.0+', 'Numerical computing for array operations'),
    ('TextBlob', '0.18.0+', 'NLP sentiment analysis (polarity scoring)'),
    ('Uvicorn', '0.30.0+', 'ASGI server for production deployment'),
]
for row_data in be_stack:
    add_table_row(table, row_data)

doc.add_paragraph()
doc.add_heading('Infrastructure', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Service'
hdr[1].text = 'Purpose'
for r in hdr:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

infra = [
    ('Render (Frontend)', 'Static site hosting for Vite build output (emenu-smart-tags-ui.onrender.com)'),
    ('Render (Backend)', 'Web service running FastAPI + Uvicorn (smart-tags-predictor.onrender.com)'),
    ('Git / GitHub', 'Version control and CI/CD integration'),
    ('localStorage', 'Client-side persistence for analysis history (no server DB needed)'),
]
for row_data in infra:
    add_table_row(table, row_data)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. System Architecture', level=1)
doc.add_paragraph(
    'The application follows a decoupled frontend-backend architecture:'
)
doc.add_paragraph()
add_code_block(
    '┌─────────────────────────────────────────────────────────────┐\n'
    '│                    FRONTEND (React + Vite)                  │\n'
    '│  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐      │\n'
    '│  │Dashboard  │ │ Analyze  │ │  Tables  │ │ History  │      │\n'
    '│  └────┬─────┘ └────┬─────┘ └────┬─────┘ └────┬─────┘      │\n'
    '│       │             │            │             │             │\n'
    '│       └─────────────┴────────────┴─────────────┘             │\n'
    '│                        api.ts                                │\n'
    '│                    (fetch wrapper)                           │\n'
    '└────────────────────────┬────────────────────────────────────┘\n'
    '                         │ HTTPS / JSON\n'
    '                         ▼\n'
    '┌─────────────────────────────────────────────────────────────┐\n'
    '│                  BACKEND (FastAPI + Uvicorn)                │\n'
    '│  ┌──────────────────────────────────────────────────────┐  │\n'
    '│  │                    api/index.py                       │  │\n'
    '│  │  /predict-guest-behavior  /predict-batch  /health    │  │\n'
    '│  │  /analyze-tags  /demo-scenarios  /simulate           │  │\n'
    '│  └──────────────────────┬───────────────────────────────┘  │\n'
    '│                         │                                   │\n'
    '│  ┌──────────────────────▼───────────────────────────────┐  │\n'
    '│  │               ml_service/predictor.py                 │  │\n'
    '│  │  GuestBehaviorPredictor (orchestrator)                │  │\n'
    '│  │    ├─ data_mapper.py  (Domain Adapter)                │  │\n'
    '│  │    ├─ model_loader.py (Keras model + preprocessor)    │  │\n'
    '│  │    ├─ sentiment.py    (TextBlob NLP)                  │  │\n'
    '│  │    └─ Smart tag rules (24 keyword rules, 6 cats)      │  │\n'
    '│  └──────────────────────┬───────────────────────────────┘  │\n'
    '│                         │                                   │\n'
    '│  ┌──────────────────────▼───────────────────────────────┐  │\n'
    '│  │           ml_service/model/fds_model_1.keras          │  │\n'
    '│  │     Sequential ANN: 128→64→32→1 (SiLU + Sigmoid)     │  │\n'
    '│  │              13,953 parameters                        │  │\n'
    '│  └──────────────────────────────────────────────────────┘  │\n'
    '└─────────────────────────────────────────────────────────────┘\n'
)

doc.add_paragraph(
    'Key architectural decisions:'
)
add_bullet('Lazy-loaded model: ', 'The Keras model loads on first prediction request, not at startup. ')
add_bullet('This avoids cold-start penalty during health checks and reduces memory until needed.')
add_bullet('Tenant isolation: ', 'Every API request carries an X-Tenant-ID header (default: restaurant_001).')
add_bullet('Blended scoring: ', 'The ANN contributes 20% and a restaurant-tuned heuristic 80% of the final score.')
add_bullet('Client-side history: ', 'Analysis results persist in localStorage, no server database required.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. MACHINE LEARNING MODEL
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Machine Learning Model', level=1)

doc.add_heading('4.1 ANN Architecture', level=2)
doc.add_paragraph(
    'The core model is a Sequential Artificial Neural Network trained on the Hotel Reservations Dataset '
    '(~36,000 records from Kaggle). It predicts the probability of a guest showing up (not canceling).'
)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Units'
hdr[2].text = 'Activation'
hdr[3].text = 'Purpose'
for r in hdr:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

layers = [
    ('Input', '27', '—', '14 numerical + 13 one-hot encoded categorical features'),
    ('Dense 1', '128', 'SiLU', 'Main feature extraction layer'),
    ('Dropout', '0.4', '—', 'Regularization to prevent overfitting'),
    ('Dense 2', '64', 'SiLU', 'Secondary feature extraction'),
    ('Dropout', '0.4', '—', 'Regularization'),
    ('Dense 3', '32', 'SiLU', 'Final hidden layer'),
    ('Dropout', '0.4', '—', 'Regularization'),
    ('Output', '1', 'Sigmoid', 'P(not_canceled) — probability of showing up'),
]
for row_data in layers:
    add_table_row(table, row_data)

doc.add_paragraph()
doc.add_paragraph('Total Parameters: 13,953')
doc.add_paragraph('Test Accuracy: 87.73%')
doc.add_paragraph('Activation Function: SiLU (Sigmoid Linear Unit) — smooth, non-monotonic activation that outperforms ReLU for this task.')

doc.add_paragraph()
doc.add_heading('Input Features (27 total)', level=3)
doc.add_paragraph('14 Numerical features (StandardScaler normalized):')
add_bullet('no_of_adults, no_of_children, no_of_weekend_nights, no_of_week_nights')
add_bullet('lead_time, arrival_year, arrival_month, arrival_date')
add_bullet('repeated_guest, no_of_previous_cancellations')
add_bullet('no_of_previous_bookings_not_canceled, avg_price_per_room')
add_bullet('required_car_parking_space, no_of_special_requests')

doc.add_paragraph('3 Categorical features (OneHotEncoder, drop="first" → 13 columns):')
add_bullet('type_of_meal_plan — 4 categories (3 encoded columns)')
add_bullet('room_type_reserved — 7 categories (6 encoded columns)')
add_bullet('market_segment_type — 5 categories (4 encoded columns)')

doc.add_paragraph()
doc.add_heading('4.2 Domain Adaptation Layer', level=2)
doc.add_paragraph(
    'Since the model was trained on hotel data but is used for restaurant predictions, a Domain Adapter '
    'bridges the gap. Restaurant inputs have fundamentally different ranges:'
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'Hotel Range'
hdr[2].text = 'Restaurant Range'
for r in hdr:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

ranges = [
    ('Lead Time', '0-443 days (mean ~85)', '0-30 days (often same-day)'),
    ('Price', '$0-$540/room (mean ~$103)', '$20-$250/cover (avg ~$60)'),
]
for row_data in ranges:
    add_table_row(table, row_data)

doc.add_paragraph()
doc.add_paragraph('adapt_lead_time() — Piecewise scaling:')
add_bullet('Same-day (0 days) → 0 (very short hotel lead)')
add_bullet('1 day → 15')
add_bullet('2-3 days → 15-65 (linear interpolation)')
add_bullet('4-7 days → 65-90')
add_bullet('8-14 days → 90-160')
add_bullet('15-30 days → 160-200')
add_bullet('30+ days → 200-350 (capped)')

doc.add_paragraph('adapt_price() — Multiplicative scaling:')
add_bullet('Under $80: multiply by 1.5x (e.g., $40 restaurant → $60 hotel)')
add_bullet('$80-$150: taper from 1.5x to 1.2x')
add_bullet('$150+: multiply by 1.2x (already in hotel range)')

doc.add_paragraph()
doc.add_paragraph('Additional mappings:')
add_bullet('Booking channel → market_segment_type: Online/App → "Online", Phone/Walk-in → "Offline", Corporate → "Corporate"')
add_bullet('Spend level → room_type: $200+ → Room_Type 4 (Premium), $120+ → Type 2, $60+ → Type 1, <$60 → Type 6 (Budget)')
add_bullet('Special needs → meal_plan: 3+ → Meal Plan 3, 1+ → Meal Plan 1, 0 → Not Selected')
add_bullet('Arrival year always set to 2018 (training era) so StandardScaler produces valid values')

doc.add_paragraph()
doc.add_heading('4.3 Blended Scoring System', level=2)
doc.add_paragraph(
    'The hotel ANN returns reliability ≈ 0.99 for almost all restaurant inputs because hotel-domain '
    'features don\'t capture restaurant-specific risk signals well. To fix this, the final score blends '
    'the ANN with a restaurant-tuned heuristic:'
)
add_code_block(
    'reliability_score = 0.20 × ANN_reliability + 0.80 × heuristic_reliability\n'
    'no_show_risk = 1.0 - reliability_score\n'
    'risk_score_display = round(no_show_risk × 100)  # 0-100%'
)
doc.add_paragraph(
    'The 20/80 weight was calibrated through iterative testing to ensure that a guest with 5 cancellations '
    'and a same-day walk-in correctly shows ~70% High Risk, while a VIP repeat guest shows ~4% Low Risk.'
)

doc.add_paragraph()
doc.add_heading('4.4 Heuristic Reliability Engine', level=2)
doc.add_paragraph(
    'The heuristic starts at a baseline of 0.65 and applies additive/subtractive adjustments:'
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Factor'
hdr[1].text = 'Adjustment'
hdr[2].text = 'Logic'
for r in hdr:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

heur = [
    ('Repeat guest', '+0.10', 'Returning guests are more reliable'),
    ('5+ completions', '+0.12', 'Strong visit history = very reliable'),
    ('3-4 completions', '+0.08', 'Moderate loyalty signal'),
    ('1-2 completions', '+0.03', 'Some track record'),
    ('5+ cancellations', '-0.45', 'Serial no-show — strongest penalty'),
    ('3-4 cancellations', '-0.30', 'High cancel history'),
    ('2 cancellations', '-0.20', 'Moderate cancel risk'),
    ('1 cancellation', '-0.10', 'Single past cancel'),
    ('30+ day advance', '-0.10', 'Very long advance = higher flake risk'),
    ('14-29 day advance', '-0.05', 'Moderate advance booking risk'),
    ('Same-day (0 days)', '-0.03', 'Impulsive booking, slightly risky'),
    ('Spend ≥ $150', '+0.08', 'High commitment through spend'),
    ('Spend ≥ $80', '+0.03', 'Moderate commitment'),
    ('Spend < $40', '-0.05', 'Low spend = less commitment'),
    ('Party size ≥ 8', '-0.05', 'Large groups harder to coordinate'),
    ('Party size ≥ 6', '-0.02', 'Moderately large group'),
]
for row_data in heur:
    add_table_row(table, row_data)

doc.add_paragraph()
doc.add_paragraph('Final score clamped to range [0.05, 0.98] to avoid 0% or 100% extremes.')

doc.add_paragraph()
doc.add_heading('4.5 Confidence Calculation', level=2)
doc.add_paragraph(
    'Confidence reflects how much the ANN and heuristic agree:'
)
add_code_block(
    'agreement = 1.0 - abs(ann_reliability - heuristic_reliability)\n'
    'confidence = 0.5 + agreement × 0.4\n'
    '# Range: 0.5 (complete disagreement) to 0.9 (perfect agreement)'
)
doc.add_paragraph(
    'If only the heuristic is available (model not loaded), confidence defaults to 0.55.'
)

doc.add_paragraph()
doc.add_heading('Risk Thresholds', level=3)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Risk Level'
hdr[1].text = 'No-Show Risk'
hdr[2].text = 'Color'
for r in hdr:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

thresholds = [
    ('High Risk', '≥ 70%', 'Red (#ef4444)'),
    ('Medium Risk', '40% - 69%', 'Amber (#f59e0b)'),
    ('Low Risk', '< 40%', 'Green (#22c55e)'),
]
for row_data in thresholds:
    add_table_row(table, row_data)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. BACKEND API
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Backend API', level=1)
doc.add_paragraph(
    'The backend is a FastAPI application (api/index.py) that serves prediction, '
    'tagging, simulation, and demo endpoints. All endpoints use JSON and support '
    'CORS from any origin.'
)

doc.add_heading('5.1 Endpoints', level=2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Method'
hdr[1].text = 'Endpoint'
hdr[2].text = 'Purpose'
hdr[3].text = 'Auth'
for r in hdr:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

endpoints = [
    ('GET', '/api/health', 'Health check, model status, version', 'None'),
    ('POST', '/api/v1/predict-guest-behavior', 'Single guest prediction with full analysis', 'X-Tenant-ID'),
    ('POST', '/api/v1/predict-batch', 'Batch prediction for multiple guests', 'X-Tenant-ID'),
    ('POST', '/api/v1/reservations/analyze-tags', 'CRM tag extraction + sentiment', 'X-Tenant-ID'),
    ('GET', '/api/v1/demo-scenarios', 'Return 18 pre-built demo scenarios', 'None'),
    ('GET', '/api/v1/simulate-reservations', 'Generate synthetic tonight\'s reservations', 'X-Tenant-ID'),
    ('GET', '/api/v1/analysis-history', 'Placeholder (history is client-side)', 'None'),
]
for row_data in endpoints:
    add_table_row(table, row_data)

doc.add_paragraph()
doc.add_heading('5.2 Request/Response Schemas', level=2)

doc.add_paragraph('ReservationInput (POST body):')
add_bullet('guest_name: string (required, 1-200 chars)')
add_bullet('party_size: int (1-20, default 2)')
add_bullet('children: int (0-10, default 0)')
add_bullet('booking_advance_days: int (≥0, default 0)')
add_bullet('special_needs_count: int (≥0, default 0)')
add_bullet('is_repeat_guest: boolean (default false)')
add_bullet('estimated_spend_per_cover: float (≥0, default 80.0)')
add_bullet('reservation_date: optional string')
add_bullet('reservation_time: optional string')
add_bullet('previous_cancellations: int (≥0, default 0)')
add_bullet('previous_completions: int (≥0, default 0)')
add_bullet('booking_channel: string (default "Online")')
add_bullet('notes: string (default "")')

doc.add_paragraph()
doc.add_paragraph('PredictionResponse:')
add_bullet('guest_name, reservation_id, tenant_id, predicted_at')
add_bullet('ai_prediction: { risk_score (0-100 int), risk_label, explanation }')
add_bullet('smart_tags: [ { category, label, color, matched } ]')
add_bullet('reliability_score, no_show_risk, risk_label, ai_tag, spend_tag')
add_bullet('sentiment: { score (0-1), label, emoji }')
add_bullet('confidence: float (0.5-0.9)')
add_bullet('explanation: string')

doc.add_paragraph()
doc.add_heading('5.3 Smart Tag Extraction', level=2)
doc.add_paragraph(
    'The predictor scans reservation notes against 24 keyword rules across 6 categories:'
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Category'
hdr[1].text = 'Tags'
hdr[2].text = 'Example Keywords'
for r in hdr:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

tags = [
    ('Dietary (9 rules)', 'Vegan, Vegetarian, Gluten Free, Dairy Free, Nut Allergy, Allergy Alert, Halal, Kosher, Jain', 'vegan, gluten-free, nut allergy, halal, kosher, jain'),
    ('Occasion (6 rules)', 'Birthday, Anniversary, Celebration, Date Night, Honeymoon, Proposal', 'birthday, anniversary, romantic, honeymoon, proposal'),
    ('Seating (4 rules)', 'Window Seat, Quiet Area, Booth, Outdoor', 'window seat, quiet, private, booth, terrace'),
    ('Status (2 rules)', 'VIP, Celebrity', 'vip, important, celebrity, famous'),
    ('Accessibility (1 rule)', 'Accessibility', 'wheelchair, accessible, disability'),
    ('Family (1 rule)', 'Family Needs', 'high chair, toddler, baby, infant'),
]
for row_data in tags:
    add_table_row(table, row_data)

doc.add_paragraph()
doc.add_paragraph(
    'Each tag includes the matched keyword for transparency, and duplicate labels are prevented '
    'via a seen_labels set. Tags are displayed with category-specific emoji icons on the frontend.'
)

doc.add_paragraph()
doc.add_heading('5.4 Sentiment Analysis', level=2)
doc.add_paragraph(
    'Uses TextBlob\'s polarity scorer on reservation notes:'
)
add_bullet('TextBlob polarity range: -1.0 to +1.0')
add_bullet('Normalized to 0-1: score = (polarity + 1.0) / 2.0')
add_bullet('Labels: ≥ 0.65 → Positive (green), ≤ 0.35 → Negative (red), else → Neutral (yellow)')
add_bullet('Empty notes default to 0.5 (Neutral)')

doc.add_paragraph()
doc.add_heading('5.5 Demo Scenarios (18 total)', level=2)
doc.add_paragraph('The system ships with 18 comprehensive demo scenarios:')

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Category'
hdr[1].text = 'Scenario Name'
hdr[2].text = 'Expected Risk'
hdr[3].text = 'Key Features'
for r in hdr:
    for p in r.paragraphs:
        for run in p.runs:
            run.bold = True

demos_data = [
    ('HIGH RISK', 'Serial No-Show', '~72%', '5 cancellations, walk-in, $35/cover, party of 6'),
    ('HIGH RISK', 'Ghost Booker', '~62%', '4 cancellations, 30-day advance, $40, party of 8'),
    ('MEDIUM', 'Risky Walk-in', '~52%', '3 cancellations, walk-in, $40, negative notes'),
    ('MEDIUM', 'Large Group Unknown', '~44%', 'Party of 10 + 3 kids, 1 cancel, 14-day advance'),
    ('LOW RISK', 'VIP Anniversary', '~8%', 'Repeat, 8 completions, $220, VIP + window seat'),
    ('LOW RISK', 'Loyal Regular', '~4%', 'Repeat, 12 completions, $180, corporate'),
    ('OCCASION', 'Birthday Party', 'Low', 'Party of 8, birthday + gluten-free tags'),
    ('OCCASION', 'Proposal Night', 'Low', '$250/cover, proposal + window seat tags'),
    ('OCCASION', 'Honeymoon Dinner', 'Low', '$200/cover, honeymoon + terrace tag'),
    ('DIETARY', 'Severe Allergy', 'Low', 'Nut allergy + vegetarian + family needs tags'),
    ('DIETARY', 'Vegan + Kosher', 'Low', 'Repeat, vegan + kosher + dairy-free tags'),
    ('DIETARY', 'Halal + Jain', 'Low', 'Halal + jain + booth tags'),
    ('SEATING', 'Window + Quiet', 'Low', 'Repeat, 6 completions, date night + seating tags'),
    ('SEATING', 'Terrace Booth', 'Low', 'Terrace + booth + celebration tags'),
    ('FAMILY', 'Family with Toddlers', 'Low', '3 kids, toddler + dairy-free tags'),
    ('ACCESS', 'Wheelchair Access', 'Low', 'Wheelchair + accessibility tags'),
    ('STATUS', 'Celebrity Guest', 'Low', '$300, celebrity + VIP + vegan + booth tags'),
    ('EDGE', 'First Timer', 'Low', 'Zero history, standard profile'),
]
for row_data in demos_data:
    add_table_row(table, row_data)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. FRONTEND APPLICATION
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Frontend Application', level=1)

doc.add_heading('6.1 Pages Overview', level=2)
doc.add_paragraph('The app uses React Router with a nested layout structure:')
add_code_block(
    'BrowserRouter\n'
    '  └─ Layout (sidebar + mobile nav + animated outlet)\n'
    '     ├─ /           → DashboardPage\n'
    '     ├─ /analyze    → AnalyzePage\n'
    '     ├─ /tables     → TableViewPage\n'
    '     ├─ /history    → HistoryPage\n'
    '     └─ /settings   → SettingsPage'
)

doc.add_paragraph()
doc.add_heading('6.2 Dashboard Page (DashboardPage.tsx)', level=2)
doc.add_paragraph(
    'The landing page uses a Bento Grid layout (12-column CSS Grid) with staggered entrance animations.'
)
doc.add_paragraph('Layout:')
add_bullet('Row 1 — 4 stat cards (3 cols each): AI Model status, Engine version, Accuracy, Domain Adapter')
add_bullet('Row 2 — Risk Gauge (5 cols) + Guest Details panel (7 cols)')
add_bullet('Row 3 — Demo Scenarios grid (8 cols, 3×N button grid) + Voice Command (4 cols)')
add_bullet('Row 4 — Navigation cards: "Analyze a Reservation" + "Table View"')
add_bullet('Bottom — API status bar showing version and domain adapter status')

doc.add_paragraph('Logic:')
add_bullet('On mount: calls healthCheck() and getDemoScenarios() in parallel')
add_bullet('Demo buttons trigger predictGuestBehavior() and update the Risk Gauge + details panel')
add_bullet('Voice Command also triggers predictions and shows results inline')

doc.add_paragraph()
doc.add_heading('6.3 Analyze Page (AnalyzePage.tsx)', level=2)
doc.add_paragraph(
    'The full-featured prediction form with a 3-column layout:'
)
add_bullet('Left (5 cols) — Input form with NumberStepper components, ChannelSelector, toggle, textarea')
add_bullet('Center (3 cols) — SimulatorControls (lead time + spend sliders) + RiskGauge')
add_bullet('Right (4 cols) — Results panel showing guest name, AI tag, spend badge, sentiment, explanation, smart tags, confidence, and SmartActions')

doc.add_paragraph('Key features:')
add_bullet('NumberStepper: Animated +/- buttons with spring pop transitions, color-coded per field (indigo for party size, violet for children, amber for special requests, red for cancellations, emerald for completions)')
add_bullet('ChannelSelector: Emoji chip buttons replacing traditional dropdown (🌐 Online, 📞 Phone, 🚶 Walk-in, 📱 App, 🏢 Corporate)')
add_bullet('Demo loading: "Load Demos" button fetches 18 scenarios; clicking one auto-fills the form')
add_bullet('Voice Input: Toggle shows VoiceCommand panel that uses Web Speech API')
add_bullet('Time-Travel Simulator: Drag lead-time or spend sliders → auto-triggers prediction with 600ms debounce')
add_bullet('History saving: Every successful prediction is saved to localStorage via saveAnalysis()')

doc.add_paragraph()
doc.add_heading('6.4 Table View Page (TableViewPage.tsx)', level=2)
doc.add_paragraph(
    'Shows tonight\'s dinner service with simulated reservations:'
)
add_bullet('"Load Tonight" button calls /simulate-reservations?count=15')
add_bullet('Reservations dated today with dinner time slots (6pm-9:30pm, weighted toward 7-8pm prime time)')
add_bullet('Summary stats bar: Total Guests, Analyzed count, High Risk count, Watch List count')
add_bullet('"Predict All" batch button calls /predict-batch for all guests at once')
add_bullet('Each row shows: risk dot (colored), guest name, REPEAT badge, notes preview, table number, time, party size, risk % badge')
add_bullet('Right panel (5 cols): RiskGauge + GuestInsightCard + SmartActions for selected guest')
add_bullet('Mobile: Bottom drawer (Vaul) shows guest details')
add_bullet('All predictions saved to history')

doc.add_paragraph()
doc.add_heading('6.5 History Page (HistoryPage.tsx)', level=2)
doc.add_paragraph(
    'Displays all past predictions stored in localStorage:'
)
add_bullet('Expandable accordion rows showing guest name, timestamp, source (Analyze/Tables), risk badge')
add_bullet('Expanded view: 3-column grid with Input Details, RiskGauge, and full Results')
add_bullet('Input details show: party size, channel, advance days, spend, cancellations, completions, repeat status, notes')
add_bullet('Results show: AI tag, spend badge, sentiment badge, explanation, smart tags, confidence bar')
add_bullet('Filter chips: All / High / Medium / Low (with counts)')
add_bullet('Individual record deletion + "Clear All" button')
add_bullet('Maximum 100 records stored (oldest auto-pruned)')

doc.add_paragraph()
doc.add_heading('6.6 Settings Page (SettingsPage.tsx)', level=2)
doc.add_paragraph('Displays system status:')
add_bullet('API Status: health check results (status, version, model loaded, service name)')
add_bullet('ML Model info: Architecture (128-64-32-1), Activation (SiLU + Sigmoid), Accuracy (87.73%), Parameters (13,953)')
add_bullet('Tenant Isolation: explains X-Tenant-ID header requirement')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. UI COMPONENT LIBRARY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('7. UI Component Library', level=1)

doc.add_heading('7.1 RiskGauge (RiskGauge.tsx)', level=2)
doc.add_paragraph(
    'A custom SVG gauge component that displays risk score from 0-100:'
)
add_bullet('270-degree arc (3/4 circle) with tick marks at 0, 25, 50, 75, 100')
add_bullet('Animated fill using requestAnimationFrame with ease-out cubic easing over 1200ms')
add_bullet('Color interpolation: Green (<40%) → Amber (40-70%) → Red (>70%)')
add_bullet('Gaussian blur glow filter on the filled arc')
add_bullet('Outer glow ring with pulsing animation for High Risk (>70%)')
add_bullet('Center display: large animated score number + "Risk Score" label')
add_bullet('Below: risk label badge with matching background color')
add_bullet('Spring entrance animation via Framer Motion')

doc.add_paragraph()
doc.add_heading('7.2 SmartTagBadge System (SmartTagBadge.tsx)', level=2)
doc.add_paragraph('A collection of badge components for displaying prediction results:')
add_bullet('SmartTagBadge: Legacy CRM tag badge (gold/blue/red/green/gray/purple themed)')
add_bullet('NoteSmartTag: Enhanced tag with category emoji (🥬 Dietary, 🎉 Occasion, 🪑 Seating, ⭐ Status, ♿ Accessibility, 👶 Family)')
add_bullet('AITagBadge: Pulsing dot + label badge for AI tags (Low Risk/green, High Spend/amber, Likely No-Show/red, Loyal Regular/indigo, Watch List/orange)')
add_bullet('SpendBadge: Tier badge (Luxury/violet, Premium/indigo, Standard/slate, Budget/muted)')
add_bullet('SentimentBadge: Emoji + label + percentage (🟢 Positive, 🟡 Neutral, 🔴 Negative)')
add_bullet('ConfidenceMeter: Horizontal progress bar with percentage (green ≥70%, amber ≥40%, red <40%)')

doc.add_paragraph()
doc.add_heading('7.3 SmartActions (SmartActions.tsx)', level=2)
doc.add_paragraph(
    'Derives actionable recommendations from prediction results:'
)
add_bullet('High Risk → "Request Deposit ($50)" (red gradient)')
add_bullet('Occasion tags → "Alert Pastry Chef" for birthday, "Prep Special Setup" for others (purple)')
add_bullet('Seating tags → "Assign [Seat Type]" (blue)')
add_bullet('Dietary/Allergy tags → "Alert Kitchen: Allergy" or "Send Dietary Card" (orange/green)')
add_bullet('Status tags → "VIP Protocol" (gold)')
add_bullet('Family tags → "Prep High Chair" (pink)')
add_bullet('Fallback → "Standard Service" (gray)')
add_bullet('Gradient buttons with shadow glow, staggered entrance animation')

doc.add_paragraph()
doc.add_heading('7.4 NumberStepper & ChannelSelector (NumberStepper.tsx)', level=2)
doc.add_paragraph('NumberStepper:')
add_bullet('Minus button — animated value display — Plus button layout')
add_bullet('Spring pop transition (damping: 20, stiffness: 300) when value changes')
add_bullet('Per-field accent colors with ring highlight and background tint')
add_bullet('Disabled states when at min/max with opacity reduction')

doc.add_paragraph('ChannelSelector:')
add_bullet('Horizontal chip button group replacing dropdown')
add_bullet('5 options: 🌐 Online, 📞 Phone, 🚶 Walk-in, 📱 App, 🏢 Corporate')
add_bullet('Active state: indigo tint with ring and shadow; inactive: subtle border')

doc.add_paragraph()
doc.add_heading('7.5 SimulatorControls (SimulatorControls.tsx)', level=2)
doc.add_paragraph(
    'The "Time-Travel Simulator" provides real-time what-if analysis:'
)
add_bullet('Lead Time slider: 0-30 days with tick labels (Now, 1d, 3d, 1w, 2w, 1mo)')
add_bullet('Spend slider: $20-$300 with $5 steps')
add_bullet('600ms debounce: changing either slider triggers automatic re-prediction')
add_bullet('AbortController: cancels in-flight requests when slider moves again')
add_bullet('Loading spinner shows "Recalculating..." during prediction')

doc.add_paragraph()
doc.add_heading('7.6 VoiceCommand (VoiceCommand.tsx)', level=2)
doc.add_paragraph(
    'Real Web Speech API integration for voice-based input:'
)
add_bullet('Uses window.SpeechRecognition or webkitSpeechRecognition (Chrome/Edge)')
add_bullet('Pulsing red rings when listening (Framer Motion infinite animations)')
add_bullet('Real-time transcript display with interim results')
add_bullet('On speech end: sets notes to transcript, triggers prediction with defaults')
add_bullet('Error handling: no-speech, not-allowed (mic denied), generic errors')
add_bullet('Graceful fallback: "Not supported in this browser" message')

doc.add_paragraph()
doc.add_heading('7.7 GuestInsightCard (GuestInsightCard.tsx)', level=2)
doc.add_paragraph('Two display modes:')
add_bullet('Compact mode: Single row with avatar, name, AI tag, reliability %, spend badge, sentiment')
add_bullet('Full mode: Multi-section card with header, 2×2 score grid (Reliability + No-Show Risk), explanation panel, smart tags, status badges, confidence meter, timestamp')
add_bullet('Risk-based glow effects: red glow for high risk, amber for medium, green for low')

doc.add_paragraph()
doc.add_heading('7.8 GuestDetailView (GuestDetailView.tsx)', level=2)
doc.add_paragraph('Responsive detail modal:')
add_bullet('Desktop (≥768px): Centered modal with backdrop blur, spring entrance animation')
add_bullet('Mobile (<768px): Vaul bottom drawer with drag handle, slides up from bottom')
add_bullet('Both use the same DetailContent component showing all prediction details')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. DATA FLOW & PREDICTION PIPELINE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Data Flow & Prediction Pipeline', level=1)
doc.add_paragraph(
    'When a user submits a reservation for analysis, the following pipeline executes:'
)

doc.add_paragraph('Step 1: Frontend Form Submission')
add_bullet('User fills form fields or selects a demo scenario')
add_bullet('Frontend calls predictGuestBehavior() which POSTs to /api/v1/predict-guest-behavior')
add_bullet('X-Tenant-ID header is included for tenant isolation')

doc.add_paragraph('Step 2: Backend Receives Request')
add_bullet('FastAPI validates the request body against ReservationInput schema (Pydantic)')
add_bullet('Lazy-loads the predictor on first call (GuestBehaviorPredictor)')

doc.add_paragraph('Step 3: Heuristic Score Computation')
add_bullet('_heuristic_reliability() computes a restaurant-tuned score from 0.05 to 0.98')
add_bullet('Applies all additive/subtractive factors (see Section 4.4)')

doc.add_paragraph('Step 4: ANN Model Inference (if available)')
add_bullet('RestaurantToHotelMapper.map_reservation() applies Domain Adapter')
add_bullet('Creates a 17-column DataFrame (14 numerical + 3 categorical)')
add_bullet('ColumnTransformer applies StandardScaler + OneHotEncoder → 27 features')
add_bullet('Keras model predicts P(not_canceled) → single float [0, 1]')

doc.add_paragraph('Step 5: Score Blending')
add_bullet('reliability = 0.20 × ANN + 0.80 × heuristic')
add_bullet('no_show_risk = 1.0 - reliability')
add_bullet('confidence = 0.5 + (1.0 - |ANN - heuristic|) × 0.4')

doc.add_paragraph('Step 6: Classification & Tag Generation')
add_bullet('Risk label assigned based on thresholds (High ≥70%, Medium ≥40%, Low <40%)')
add_bullet('AI tag: "Likely No-Show", "High Spend Potential", "Loyal Regular", "Watch List", or "Low Risk"')
add_bullet('Spend tier: Luxury ($200+), Premium ($120+), Standard ($60+), Budget')
add_bullet('Smart tags extracted from notes via 24 keyword rules')
add_bullet('Sentiment analyzed via TextBlob polarity')
add_bullet('Explanation generated from top 3 factors')

doc.add_paragraph('Step 7: Response & Frontend Display')
add_bullet('API returns PredictionResponse with all computed fields')
add_bullet('Frontend updates RiskGauge, badges, explanation, tags, actions')
add_bullet('Result saved to localStorage history')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 9. STYLING & DESIGN SYSTEM
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Styling & Design System', level=1)

doc.add_heading('Design Philosophy: Dark Glassmorphism', level=2)
doc.add_paragraph(
    'The UI uses a dark glassmorphism design language inspired by modern SaaS dashboards:'
)
add_bullet('Base background: slate-950 (#020617) — near-black blue')
add_bullet('Glass cards: rgba(255,255,255,0.05) background with 24px backdrop-blur')
add_bullet('Glass borders: rgba(255,255,255,0.1) — subtle white borders')
add_bullet('Hover state: rgba(255,255,255,0.08) — slightly brighter on interaction')
add_bullet('Text hierarchy: white (primary), slate-400 (secondary), slate-500 (labels), slate-600 (muted)')
add_bullet('Accent color: Indigo (#6366f1) — used for primary buttons, active states, links')

doc.add_paragraph()
doc.add_heading('CSS Custom Properties', level=3)
add_code_block(
    '--color-glass: rgba(255, 255, 255, 0.05)\n'
    '--color-glass-border: rgba(255, 255, 255, 0.1)\n'
    '--color-glass-hover: rgba(255, 255, 255, 0.08)'
)

doc.add_heading('Component Classes', level=3)
add_bullet('.glass — Glass card with backdrop-blur, border, rounded-2xl')
add_bullet('.glass-hover — Hover transition: brighter background + border')
add_bullet('.btn-primary — Indigo-600, white text, rounded-xl, active:scale-95')
add_bullet('.btn-ghost — Transparent with slate-400 text, hover:white + bg')
add_bullet('.input-dark — Dark input with white/5 bg, indigo focus ring')
add_bullet('.glow-red/amber/green/indigo — Box-shadow glow effects for risk levels')
add_bullet('.animate-pulse-red — Pulsing red glow for high-risk elements')
add_bullet('.slider-cyber — Custom range slider with indigo thumb and glow effect')

doc.add_heading('Animation Patterns', level=3)
add_bullet('Page transitions: opacity + y-offset with 200ms duration (AnimatePresence mode="wait")')
add_bullet('Staggered entrance: parent container with staggerChildren: 0.03-0.08s')
add_bullet('Number changes: spring pop (damping: 20, stiffness: 300) with y-axis slide')
add_bullet('Button interactions: whileHover scale 1.02-1.05, whileTap scale 0.85-0.97')
add_bullet('Risk gauge: 1200ms ease-out cubic animation from 0 to target')
add_bullet('Voice listening: infinite scale 1→2.5 pulse rings with red-500/30')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 10. DEPLOYMENT
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Deployment & Infrastructure', level=1)

doc.add_heading('Backend Deployment (Render Web Service)', level=2)
add_bullet('Runtime: Python 3.11')
add_bullet('Build command: ./build.sh (installs pip deps + downloads NLTK data)')
add_bullet('Start command: uvicorn api.index:app --host 0.0.0.0 --port $PORT')
add_bullet('URL: https://smart-tags-predictor.onrender.com')
add_bullet('Cold start: ~30-60s (TensorFlow + model loading)')
add_bullet('Model file: ml_service/model/fds_model_1.keras (~170KB)')
add_bullet('Preprocessor: Cached as preprocessor.pkl after first fit')

doc.add_heading('Frontend Deployment (Render Static Site)', level=2)
add_bullet('Build command: npm install && npm run build')
add_bullet('Publish directory: dist/')
add_bullet('URL: https://emenu-smart-tags-ui.onrender.com')
add_bullet('Bundle size: ~453KB JS (139KB gzipped) + 55KB CSS (9KB gzipped)')
add_bullet('SPA routing: All paths serve index.html')

doc.add_heading('Environment Configuration', level=3)
add_bullet('Frontend API_BASE: In production, points to https://smart-tags-predictor.onrender.com/api')
add_bullet('In development, Vite proxies /api to localhost:8000')
add_bullet('No environment variables required (all config is in code)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 11. FILE STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('11. File Structure', level=1)

add_code_block(
    'smart-tags-prediction-model/\n'
    '├── api/\n'
    '│   └── index.py                 # FastAPI backend (all endpoints)\n'
    '├── ml_service/\n'
    '│   ├── __init__.py\n'
    '│   ├── predictor.py             # GuestBehaviorPredictor + heuristic + tags\n'
    '│   ├── model_loader.py          # Keras model + sklearn preprocessor loader\n'
    '│   ├── data_mapper.py           # Restaurant→Hotel feature mapping + Domain Adapter\n'
    '│   ├── sentiment.py             # TextBlob sentiment analysis\n'
    '│   ├── data_simulator.py        # Synthetic reservation generator\n'
    '│   └── model/\n'
    '│       ├── fds_model_1.keras    # Trained ANN model (13,953 params)\n'
    '│       └── Hotel_reservations.csv  # Original training dataset\n'
    '├── src/\n'
    '│   ├── App.tsx                  # Router setup\n'
    '│   ├── main.tsx                 # React entry point\n'
    '│   ├── index.css                # Global styles (glassmorphism, sliders)\n'
    '│   ├── lib/\n'
    '│   │   ├── api.ts              # API client (fetch wrappers)\n'
    '│   │   ├── types.ts            # TypeScript interfaces\n'
    '│   │   └── historyStore.ts     # localStorage persistence\n'
    '│   ├── components/\n'
    '│   │   ├── Layout.tsx          # Sidebar + mobile nav + outlet\n'
    '│   │   ├── RiskGauge.tsx       # SVG gauge with animations\n'
    '│   │   ├── SmartTagBadge.tsx   # Badge component collection\n'
    '│   │   ├── SmartActions.tsx    # Action recommendation buttons\n'
    '│   │   ├── NumberStepper.tsx   # Animated +/- input + ChannelSelector\n'
    '│   │   ├── SimulatorControls.tsx  # Time-Travel sliders\n'
    '│   │   ├── VoiceCommand.tsx    # Web Speech API integration\n'
    '│   │   ├── GuestInsightCard.tsx   # Full prediction display card\n'
    '│   │   ├── GuestDetailView.tsx    # Modal (desktop) / Drawer (mobile)\n'
    '│   │   └── ResultCard.tsx      # Legacy result card\n'
    '│   └── pages/\n'
    '│       ├── DashboardPage.tsx   # Bento grid dashboard\n'
    '│       ├── AnalyzePage.tsx     # Full prediction form + simulator\n'
    '│       ├── TableViewPage.tsx   # Tonight\'s service view\n'
    '│       ├── HistoryPage.tsx     # Past analysis browser\n'
    '│       └── SettingsPage.tsx    # System status display\n'
    '├── build.sh                    # Backend build script\n'
    '├── requirements.txt            # Python dependencies\n'
    '├── package.json                # Node.js dependencies\n'
    '├── vite.config.ts              # Vite + Tailwind config\n'
    '├── tsconfig.json               # TypeScript config\n'
    '└── index.html                  # HTML entry point\n'
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Report —')
run.bold = True
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

# Save
output_path = '/home/user/smart-tags-prediction-model/eMenu_Smart_Tags_Technical_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
